import streamlit as st
import requests
import urllib.parse
import re
import io
import jieba
from pypinyin import pinyin, Style
from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Cm, Pt

st.set_page_config(page_title="APA 7 智能排版引擎", layout="wide")

st.title("APA 7 自动文献排版与 DOI 抓取工具")
st.markdown("专为跨文化/人文学科研究设计。支持**批量文本智能识别**与**精确手动微调**双模式。")
st.divider()

if "raw_text_input" not in st.session_state:
    st.session_state.raw_text_input = ""

def clear_text():
    st.session_state.raw_text_input = ""

def create_word_docx(results_list):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    for res in results_list:
        clean_res = res.replace('\n', '').replace('\r', '').strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        p.paragraph_format.line_spacing = 1.15

        parts = clean_res.split('*')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            if i % 2 != 0:
                run.italic = True
            run.font.name = 'Times New Roman'
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def get_doi_from_crossref(title, authors="", year=""):
    if not title: return ""
    query = f"{title} {authors}".strip()
    url = f"https://api.crossref.org/works?query.bibliographic={urllib.parse.quote(query)}&select=DOI,title,author,issued&rows=1"
    try:
        headers = {'User-Agent': 'AcademicCitationTool/2.0'}
        response = requests.get(url, headers=headers, timeout=5)
        if response.status_code == 200:
            data = response.json()
            items = data.get('message', {}).get('items', [])
            if items:
                item = items[0]
                doi = item.get('DOI')
                if not doi: return ""
                
                if year:
                    issued = item.get('issued', {}).get('date-parts', [[None]])
                    api_year = str(issued[0][0]) if issued and issued[0][0] else ""
                    if api_year and str(year) not in api_year: return ""
                        
                raw_lower = f"{authors} {title}".lower()
                if item.get('author'):
                    api_author = item['author'][0].get('family', '').lower()
                    if api_author and api_author not in raw_lower: return ""
                
                api_title = item.get('title', [''])[0].lower()
                api_title_clean = re.sub(r'[^a-z0-9]', '', api_title)
                if len(api_title_clean) > 4:
                    if api_title_clean[:5] not in re.sub(r'[^a-z0-9]', '', raw_lower):
                        return ""
                
                return f"https://doi.org/{doi}"
    except Exception:
        pass
    return ""

def smart_title_case(text):
    if not text: return ""
    words = text.split()
    minor_words = {'and', 'or', 'but', 'nor', 'a', 'an', 'the', 'as', 'at', 'by', 'for', 'in', 'of', 'on', 'to', 'with', 'do', 'da', 'de', 'das', 'dos'}
    res = []
    for i, w in enumerate(words):
        prefix = ""
        core = w
        if w[0] in "([{‘“'\"":
            prefix = w[0]
            core = w[1:]
            
        if not core: 
            res.append(w)
            continue

        if core.isupper() or (len(core) > 1 and core[1] == '.' and core[0].isupper()):
            if core.lower() in minor_words and i != 0: res.append(prefix + core.lower())
            else: res.append(prefix + core)
            continue
            
        lower_core = core.lower()
        if i != 0 and lower_core in minor_words: 
            res.append(prefix + lower_core)
        else:
            if len(core) > 1: res.append(prefix + core[0].upper() + core[1:].lower())
            else: res.append(prefix + core.upper())
    return " ".join(res).replace("'S", "'s")

def enforce_sentence_case(text):
    if not text: return ""
    
    # 【新增：防御全大写 API 幻觉】
    # 如果超过 40% 的字母是大写，判定为 API 传来的脏数据，强行降维处理
    letters = [c for c in text if c.isalpha()]
    if letters and sum(1 for c in letters if c.isupper()) / len(letters) > 0.4:
        text = text.lower()
        
    words = text.split()
    res = []
    for w in words:
        # 仅保护真正的缩写 (如 USA) 和混合词 (如 iPhone)
        if (w.isupper() and len(re.findall(r'[A-Z]', w)) >= 2) or re.search(r'[a-z][A-Z]', w): 
            res.append(w)
        else: 
            res.append(w.lower())
    
    s = " ".join(res)
    s = re.sub(r'^([^a-zA-Z]*)([a-zA-Z])', lambda m: m.group(1) + m.group(2).upper(), s)
    s = re.sub(r'([:\?\!]\s+)([^a-zA-Z]*)([a-zA-Z])', lambda m: m.group(1) + m.group(2) + m.group(3).upper(), s)
    s = re.sub(r'([\[\(]\s*)([a-zA-Z])', lambda m: m.group(1) + m.group(2).upper(), s)
    return s

def convert_authors_to_apa(chinese_authors):
    if not chinese_authors: return ""
    if not re.search(r'[\u4e00-\u9fff]', chinese_authors):
        return re.sub(r'\b([a-z])\b', lambda m: m.group(1).upper(), chinese_authors).strip()

    author_list = [a.strip() for a in re.split(r'[,，、]', chinese_authors) if a.strip()]
    processed_authors = []
    for author in author_list:
        py_list = pinyin(author, style=Style.NORMAL)
        if len(py_list) == 1:
            processed_authors.append(py_list[0][0].capitalize())
        else:
            surname = py_list[0][0].capitalize()
            initials = " ".join([item[0][0].upper() + "." for item in py_list[1:]])
            processed_authors.append(f"{surname}, {initials}")
            
    processed_authors = [re.sub(r'\s+\.', '.', a) for a in processed_authors]

    if len(processed_authors) == 1: return processed_authors[0]
    elif len(processed_authors) == 2: return f"{processed_authors[0]}, & {processed_authors[1]}"
    else: return ", ".join(processed_authors[:-1]) + f", & {processed_authors[-1]}"

def convert_editor_to_apa(chinese_name):
    if not chinese_name: return ""
    py_list = pinyin(chinese_name, style=Style.NORMAL)
    if len(py_list) < 2: return py_list[0][0].capitalize()
    surname = py_list[0][0].capitalize()
    given_initials = " ".join([item[0][0].upper() + "." for item in py_list[1:]])
    return f"{given_initials} {surname}".replace(" .", ".")

def convert_title_to_pinyin_grouped(chinese_title):
    if not chinese_title: return ""
    clean_title = re.sub(r'^[\.。\s]+', '', chinese_title)
    words = jieba.lcut(clean_title)
    pinyin_words = []
    for word in words:
        if not word.strip(): continue # 防止空字符变成额外空格
        word_pinyin_list = pinyin(word, style=Style.NORMAL)
        word_pinyin_str = "".join([item[0] for item in word_pinyin_list])
        if word_pinyin_str: pinyin_words.append(word_pinyin_str)
    if not pinyin_words: return ""
    
    pinyin_str = " ".join(pinyin_words).strip()
    
    # 【新增：拼音标点真空吸尘器】
    pinyin_str = re.sub(r'\s+([\)\]\.,!\?])', r'\1', pinyin_str) # 吸走右括号前、句号前的空格
    pinyin_str = re.sub(r'([\[\(])\s+', r'\1', pinyin_str)     # 吸走左括号后的空格
    
    if len(pinyin_str) > 1: return pinyin_str[0].upper() + pinyin_str[1:].lower()
    return pinyin_str.upper()

def translate_text(text, target_lang='pt'):
    if not text: return ""
    try:
        return GoogleTranslator(source='zh-CN', target=target_lang).translate(text)
    except Exception:
        return ""

def fallback_parse(raw_text, is_zh):
    empty_defaults = {'volume': '', 'issue': '', 'page_range': '', 'publisher': '', 'book_title': '', 'editors': '', 'doi': ''}
    
    year_match = re.search(r'[\(（]\s*(\d{4}[a-z]?)\s*[\)）]', raw_text)
    if not year_match:
        res = {'ref_type': 'journal', 'is_chinese': is_zh, 'author': raw_text, 'year': '202X', 'title': '', 'journal': ''}
        res.update(empty_defaults)
        return res

    year = year_match.group(1)
    parts = re.split(r'[\(（]\s*' + re.escape(year) + r'\s*[\)）]', raw_text, maxsplit=1)
    
    author = parts[0].strip()
    rest = re.sub(r'^[\.,\s]+', '', parts[1])

    url_match = re.search(r'(https?://\S+|doi:\S+)', rest, re.IGNORECASE)
    extracted_url = ""
    if url_match:
        extracted_url = url_match.group(1)
        rest = rest.replace(url_match.group(0), '').strip(' .')

    if re.search(r'(?:[\.\?!\]"”\'])\s+In\b', rest, re.IGNORECASE) or rest.lower().startswith("in "):
        title_parts = re.split(r'\.\s+In\s+|\.\s+in\s+', rest, maxsplit=1)
        if len(title_parts) == 1 and rest.lower().startswith("in "): title, in_part = "", rest[3:].strip()
        elif len(title_parts) == 2: title, in_part = title_parts[0].strip(), title_parts[1].strip()
        else: title, in_part = rest, ""

        editors, book_title, page_range, publisher, volume = "", "", "", "", ""
        if in_part:
            ed_split = re.split(r'\([E|O]d[^\)]*\),?\s*|\(Org[^\)]*\),?\s*', in_part, flags=re.IGNORECASE)
            if len(ed_split) >= 2:
                editors, book_part = ed_split[0].strip(), ed_split[-1].strip()
                page_match = re.search(r'\(\s*(?:(?:vol\.|Vol\.)\s*(\d+)\s*,?\s*)?pp\.\s*([\d\-–]+)\s*\)\.?\s*(.*)', book_part, re.IGNORECASE)
                if page_match:
                    volume = page_match.group(1) if page_match.group(1) else ""
                    page_range = page_match.group(2)
                    publisher = page_match.group(3).strip()
                    book_title = book_part[:page_match.start()].strip(" ,.")
                else: book_title = book_part
            else: book_title = in_part
                
        return {'ref_type': 'chapter', 'is_chinese': is_zh, 'author': author, 'year': year, 'title': title, 'journal': '', 'volume': volume, 'issue': '', 'page_range': page_range, 'publisher': publisher, 'book_title': book_title, 'editors': editors, 'doi': extracted_url}

    journal_pattern = re.search(r',\s*(\d+)(?:\s*[\(（]([^)）]+)[\)）])?\s*,\s*([pP]*\.?\s*[\d\-–—]+)\.?$', rest)
    if journal_pattern:
        volume = journal_pattern.group(1)
        issue = journal_pattern.group(2) if journal_pattern.group(2) else ""
        page_range = journal_pattern.group(3)
        
        prefix = rest[:journal_pattern.start()].strip()
        last_dot_idx = prefix.rfind('.')
        if last_dot_idx != -1:
            title = prefix[:last_dot_idx].strip()
            journal = prefix[last_dot_idx+1:].strip()
        else:
            title, journal = prefix, ""
            
        res = {'ref_type': 'journal', 'is_chinese': is_zh, 'author': author, 'year': year, 'title': title, 'journal': journal, 'volume': volume, 'issue': issue, 'page_range': page_range}
        res.update({'publisher': '', 'book_title': '', 'editors': '', 'doi': extracted_url})
        return res

    book_parts = [p.strip() for p in re.split(r'[\.。]', rest) if p.strip()]
    if len(book_parts) >= 2:
        publisher = book_parts[-1]
        title = ". ".join(book_parts[:-1])
    else:
        title, publisher = rest, ""
        
    res = {'ref_type': 'book', 'is_chinese': is_zh, 'author': author, 'year': year, 'title': title, 'journal': '', 'publisher': publisher}
    res.update({'volume': '', 'issue': '', 'page_range': '', 'book_title': '', 'editors': '', 'doi': extracted_url})
    return res

def parse_raw_citation_via_crossref(raw_text):
    url = f"https://api.crossref.org/works?query.bibliographic={urllib.parse.quote(raw_text)}&rows=1"
    try:
        headers = {'User-Agent': 'AcademicCitationTool/2.0'}
        response = requests.get(url, headers=headers, timeout=8)
        if response.status_code == 200:
            data = response.json()
            items = data.get('message', {}).get('items', [])
            if items:
                item = items[0]
                raw_lower = raw_text.lower()
                issued = item.get('issued', {}).get('date-parts', [[None]])
                api_year = str(issued[0][0]) if issued and issued[0][0] else ""
                if not api_year or api_year not in raw_text: return None
                if item.get('author'):
                    api_author = item['author'][0].get('family', '').lower()
                    if not api_author or api_author not in raw_lower: return None
                api_title = item.get('title', [''])[0].lower()
                api_title_clean = re.sub(r'[^a-z0-9]', '', api_title)
                if len(api_title_clean) > 4 and api_title_clean[:5] not in re.sub(r'[^a-z0-9]', '', raw_lower):
                    return None
                
                c_type = item.get('type', '')
                if c_type == 'journal-article': ref_type = 'journal'
                elif c_type == 'book-chapter': ref_type = 'chapter'
                else: ref_type = 'book'
                
                container_title = item.get('container-title', [''])[0]
                if ref_type == 'journal' and not container_title: 
                    return None
                
                authors_list = []
                for a in item.get('author', []):
                    family, given = a.get('family', '').strip(), a.get('given', '').strip()
                    if family.isupper(): family = family.title()
                    if family and given:
                        parts = given.replace('.', ' ').replace('-', ' ').split()
                        initials = " ".join([p[0].upper() + "." for p in parts if p])
                        authors_list.append(f"{family}, {initials}")
                    elif family: authors_list.append(family)
                if len(authors_list) == 1: author_str = authors_list[0]
                elif len(authors_list) == 2: author_str = f"{authors_list[0]}, & {authors_list[1]}"
                elif len(authors_list) > 2: author_str = ", ".join(authors_list[:-1]) + f", & {authors_list[-1]}"
                else: author_str = ""

                editors_list = []
                for e in item.get('editor', []):
                    family, given = e.get('family', '').strip(), e.get('given', '').strip()
                    if family.isupper(): family = family.title()
                    if family and given:
                        parts = given.replace('.', ' ').replace('-', ' ').split()
                        initials = " ".join([p[0].upper() + "." for p in parts if p])
                        editors_list.append(f"{initials} {family}")
                    elif family: editors_list.append(family)
                if len(editors_list) == 1: editor_str = editors_list[0]
                elif len(editors_list) == 2: editor_str = f"{editors_list[0]} & {editors_list[1]}"
                elif len(editors_list) > 2: editor_str = ", ".join(editors_list[:-1]) + f", & {editors_list[-1]}"
                else: editor_str = ""
                
                title_main = item.get('title', [''])[0] if item.get('title') else ''
                subtitle = item.get('subtitle', [''])[0] if item.get('subtitle') else ''
                title = f"{title_main}: {subtitle}" if subtitle else title_main
                
                api_page = str(item.get('page', '')).strip()
                if '-' not in api_page and '–' not in api_page:
                    page_match = re.search(r'\b(\d+[-–—]\d+)\b', raw_text)
                    if page_match: api_page = page_match.group(1).replace('–', '-').replace('—', '-')

                doi = item.get('DOI', '')
                if doi: doi = f"https://doi.org/{doi}".strip()
                
                return {
                    'ref_type': ref_type, 'is_chinese': False, 'author': author_str,
                    'year': api_year, 'title': title, 'journal': container_title if ref_type == 'journal' else '',
                    'book_title': container_title if ref_type == 'chapter' else '',
                    'volume': item.get('volume', ''), 'issue': item.get('issue', ''),
                    'page_range': api_page, 'publisher': item.get('publisher', ''),
                    'editors': editor_str, 'doi': doi
                }
    except Exception:
        pass
    return None

class ReferenceItem:
    def __init__(self, ref_type, author, year, title, is_chinese=True, 
                 journal="", volume="", issue="", page_range="", 
                 publisher="", book_title="", editors="", doi="", target_lang='en'):
        self.ref_type = ref_type          
        self.is_chinese = is_chinese      
        self.author = author          
        self.year = year              
        self.title = title      
        self.page_range = page_range  
        self.doi = doi                
        self.journal = journal        
        self.volume = volume          
        self.issue = issue            
        self.publisher = publisher    
        self.book_title = book_title  
        self.editors = editors 
        self.target_lang = target_lang

    def separate_edition(self, title_str):
        if not title_str: return "", ""
        m = re.search(r'^(.*?)\s*(\([^)]*(?:ed\.|Ed\.|edição|edição|vol\.|Vol\.)[^)]*\))\.?$', title_str, re.IGNORECASE)
        if m:
            return m.group(1).strip(), m.group(2).strip()
        return title_str.strip(), ""

    def to_apa_string(self):
        if not self.doi:
            self.doi = get_doi_from_crossref(self.title, self.author, self.year)

        final_author = convert_authors_to_apa(self.author) if self.is_chinese else self.author
        final_author = final_author.strip()
        final_author = re.sub(r'[\s\.]+$', '', final_author)
        if final_author: final_author += '.'

        if self.is_chinese:
            py_title = convert_title_to_pinyin_grouped(self.title)
            tr_title = translate_text(self.title, target_lang=self.target_lang)
            if tr_title: tr_title = enforce_sentence_case(tr_title)
            
            if self.ref_type == 'book':
                main_py, paren_py = self.separate_edition(py_title)
                final_title = f"*{main_py}* {paren_py} [{tr_title}]".strip() if paren_py else f"*{py_title}* [{tr_title}]"
            else:
                final_title = f"{py_title} [{tr_title}]"
        else:
            raw_title = enforce_sentence_case(self.title)
            if self.ref_type == 'book':
                main_t, paren_t = self.separate_edition(raw_title)
                final_title = f"*{main_t}* {paren_t}".strip()
            else:
                final_title = raw_title

        doi_str = f" {self.doi}" if self.doi else ""
        final_result = ""

        if self.ref_type == 'journal':
            if self.is_chinese:
                if self.journal:
                    py_journal = smart_title_case(convert_title_to_pinyin_grouped(self.journal))
                    tr_journal = smart_title_case(translate_text(self.journal, target_lang='en'))
                    final_journal = f"*{py_journal}* ({tr_journal})"
                else:
                    final_journal = ""
            else:
                final_journal = f"*{smart_title_case(self.journal)}*" if self.journal else ""
            
            vol_iss = ""
            if self.volume and self.issue:
                vol_iss = f"*{self.volume}*({self.issue})"
            elif self.volume:
                vol_iss = f"*{self.volume}*"
                
            journal_parts = []
            if final_journal: journal_parts.append(final_journal)
            if vol_iss: journal_parts.append(vol_iss)
            if self.page_range: journal_parts.append(self.page_range)
            
            journal_part_str = ", ".join(journal_parts)
            journal_suffix = f" {journal_part_str}." if journal_part_str else ""
            final_result = f"{final_author} ({self.year}). {final_title}.{journal_suffix}{doi_str}".strip()

        elif self.ref_type == 'book':
            final_publisher = smart_title_case(translate_text(self.publisher, target_lang='en')) if self.is_chinese else self.publisher
            pub_part = f" {final_publisher}." if final_publisher else ""
            final_result = f"{final_author} ({self.year}). {final_title}.{pub_part}{doi_str}".strip()

        elif self.ref_type == 'chapter':
            if self.is_chinese:
                final_editors = convert_editor_to_apa(self.editors)
                py_book = convert_title_to_pinyin_grouped(self.book_title)
                tr_book = translate_text(self.book_title, target_lang=self.target_lang)
                if tr_book: tr_book = enforce_sentence_case(tr_book)
                main_py, paren_py = self.separate_edition(py_book)
                final_book_title = f"*{main_py}* {paren_py} [{tr_book}]".strip() if paren_py else f"*{py_book}* [{tr_book}]"
                final_publisher = smart_title_case(translate_text(self.publisher, target_lang='en'))
            else:
                final_editors = self.editors
                raw_book = enforce_sentence_case(self.book_title)
                main_bt, paren_bt = self.separate_edition(raw_book)
                final_book_title = f"*{main_bt}* {paren_bt}".strip()
                final_publisher = self.publisher
            
            ed_suffix = "(Eds.)" if "&" in final_editors or "," in final_editors else "(Ed.)"
            editor_part = f"In {final_editors} {ed_suffix}, " if final_editors else ""
            
            vol_page_parts = []
            if self.volume: vol_page_parts.append(f"Vol. {self.volume}")
            if self.page_range:
                if "pp." in self.page_range.lower() or "p." in self.page_range.lower(): vol_page_parts.append(self.page_range)
                else: vol_page_parts.append(f"pp. {self.page_range}")
            
            page_str = f" ({', '.join(vol_page_parts)})" if vol_page_parts else ""
            pub_part = f" {final_publisher}." if final_publisher else ""
            final_result = f"{final_author} ({self.year}). {final_title}. {editor_part}{final_book_title}{page_str}.{pub_part}{doi_str}".strip()
        
        final_result = re.sub(r'\s+', ' ', final_result)
        final_result = final_result.replace(' .', '.')
        final_result = final_result.replace('..', '.')
        final_result = final_result.replace('**.', '')
        
        return final_result

col_lang1, col_lang2 = st.columns([1, 2])
with col_lang1:
    target_lang_choice = st.radio("翻译目标语言 / Target Language", ["英语 (English)", "葡萄牙语 (Português)"], horizontal=True)
target_lang_code = 'en' if '英语' in target_lang_choice else 'pt'

tab1, tab2 = st.tabs(["🚀 批量智能解析 (推荐)", "✍️ 手动精准录入"])

with tab1:
    st.markdown("把杂乱的、缺失 DOI 或排版不准的文献直接粘贴在下面。系统会自动提取字段、查缺补漏并重排。")
    st.text_area("在此粘贴参考文献原始文本：", height=200, key="raw_text_input")
    
    col_btn1, col_btn2 = st.columns([4, 1])
    with col_btn1:
        run_batch = st.button("一键智能整理", type="primary", use_container_width=True)
    with col_btn2:
        st.button("🗑️ 清空内容", on_click=clear_text, use_container_width=True)
    
    if run_batch:
        if not st.session_state.raw_text_input.strip():
            st.warning("⚠️ 请先粘贴文本！")
        else:
            lines = [line.strip() for line in st.session_state.raw_text_input.split('\n') if line.strip()]
            results = []
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for i, line in enumerate(lines):
                status_text.text(f"正在智能解析第 {i+1}/{len(lines)} 条文献...")
                is_zh = len(re.findall(r'[\u4e00-\u9fff]', line)) > 5
                
                parsed_data = None
                if not is_zh:
                    parsed_data = parse_raw_citation_via_crossref(line)
                    if parsed_data and parsed_data['ref_type'] == 'journal':
                        if re.search(r'\bIn\b', line, re.IGNORECASE) and (re.search(r'\(Eds?\.?\)', line, re.IGNORECASE) or re.search(r'\(Orgs?\.?\)', line, re.IGNORECASE)):
                            api_doi = parsed_data.get('doi', '')
                            parsed_data = fallback_parse(line, is_zh)
                            parsed_data['doi'] = api_doi

                if not parsed_data:
                    parsed_data = fallback_parse(line, is_zh)
                    
                ref = ReferenceItem(**parsed_data, target_lang=target_lang_code)
                
                if not is_zh and parsed_data.get('is_chinese') == False:
                    ref.author = parsed_data.get('author', '')
                    ref.title = parsed_data.get('title', '')
                    ref.journal = parsed_data.get('journal', '')
                    ref.book_title = parsed_data.get('book_title', '')
                    ref.publisher = parsed_data.get('publisher', '')
                    ref.editors = parsed_data.get('editors', '')
                
                final_str = ref.to_apa_string()
                results.append(final_str)
                progress_bar.progress((i + 1) / len(lines))
                
            status_text.text("✅ 处理完成！")
            results.sort(key=lambda x: x.lstrip('*'))
            st.subheader("👀 渲染效果预览")
            for r in results:
                st.markdown(f"- {r}")
                
            st.subheader("💾 导出文件")
            word_file = create_word_docx(results)
            st.download_button(
                label="📥 一键下载为 Word 文档 (.docx)",
                data=word_file,
                file_name="APA7_References.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )

with tab2:
    st.markdown("如果智能解析有个别字段不准，可以在这里进行精准的手动分项输入。")
    col1, col2 = st.columns(2)
    with col1:
        ref_type = st.selectbox("文献类型", ["journal", "book", "chapter"], format_func=lambda x: {"journal":"期刊文章", "book":"书籍专著", "chapter":"书籍章节"}[x])
        is_chinese = st.radio("语言类型", ["中文文献 (执行自动拼音化与翻译)", "外文文献 (保持字母原样)"]) == "中文文献 (执行自动拼音化与翻译)"
        author = st.text_input("作者 (如: 张三、李四)")
        year = st.text_input("年份")
        title = st.text_input("文章/书籍标题")
    with col2:
        journal = st.text_input("期刊名 (仅期刊填写)")
        volume = st.text_input("卷号 Volume")
        issue = st.text_input("期号 Issue")
        book_title = st.text_input("母书名 (仅书籍章节)")
        editors = st.text_input("编者 (仅书籍章节)")
        publisher = st.text_input("出版社")
        page_range = st.text_input("页码范围")
        doi = st.text_input("DOI (留空则自动抓取)")

    if st.button("生成单条引用", type="primary", use_container_width=True):
        if not author or not year or not title:
            st.error("⚠️ 错误：请至少填写作者、年份和标题！")
        else:
            with st.spinner("正在处理排版..."):
                ref = ReferenceItem(ref_type=ref_type, is_chinese=is_chinese, author=author, year=year, title=title, journal=journal, volume=volume, issue=issue, page_range=page_range, publisher=publisher, book_title=book_title, editors=editors, doi=doi, target_lang=target_lang_code)
                result = ref.to_apa_string()
                st.success("✅ 生成成功！")
                st.markdown(result)
                
                word_file_single = create_word_docx([result])
                st.download_button(
                    label="📥 下载为 Word 文档 (.docx)",
                    data=word_file_single,
                    file_name="Single_Reference.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )